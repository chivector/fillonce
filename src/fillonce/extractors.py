from __future__ import annotations

import csv
import hashlib
import json
import re
from collections.abc import Iterator
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from pypdf import PdfReader

from fillonce.models import Evidence, Fact
from fillonce.normalization import (
    canonical_key,
    display_value,
    looks_like_date_key,
    normalize_date,
    normalize_value,
)

SUPPORTED_EXTENSIONS = {".json", ".yaml", ".yml", ".csv", ".txt", ".md", ".docx", ".pdf"}
PAIR_PATTERN = re.compile(r"^\s*[-*]?\s*([^:\n]{1,80}?)\s*[:：]\s*(\S.*)\s*$")
EMAIL_PATTERN = re.compile(r"(?<![\w.+-])[\w.+-]+@(?:[\w-]+\.)+[A-Za-z]{2,}(?![\w-])")
PHONE_PATTERN = re.compile(r"(?<!\w)(?:\+?\d[\d ()-]{7,}\d)(?!\w)")
URL_PATTERN = re.compile(r"https?://[^\s)>]+")


class ExtractionError(ValueError):
    pass


def _fact_id(source: Path, locator: str, key: str, value: str) -> str:
    digest = hashlib.sha256(f"{source}|{locator}|{key}|{value}".encode()).hexdigest()[:12]
    return f"fact_{digest}"


def _make_fact(source: Path, locator: str, label: str, value: object) -> Fact | None:
    rendered = display_value(value)
    if not label.strip() or not rendered:
        return None
    key = canonical_key(label)
    if looks_like_date_key(key):
        rendered = normalize_date(rendered)
    evidence = Evidence(source=str(source), locator=locator, excerpt=f"{label}: {rendered}")
    return Fact(
        fact_id=_fact_id(source, locator, key, rendered),
        key=key,
        label=label.strip(),
        value=rendered,
        normalized_value=normalize_value(rendered),
        evidence=evidence,
    )


def _flatten(data: Any, prefix: str = "") -> Iterator[tuple[str, object, str]]:
    if isinstance(data, dict):
        for key, value in data.items():
            path = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(value, (dict, list)):
                yield from _flatten(value, path)
            elif value is not None:
                yield str(key), value, path
    elif isinstance(data, list):
        for index, value in enumerate(data):
            path = f"{prefix}[{index}]"
            if isinstance(value, (dict, list)):
                yield from _flatten(value, path)
            elif value is not None:
                label = prefix.rsplit(".", 1)[-1] if prefix else f"item_{index + 1}"
                yield label, value, path


def _extract_structured(path: Path) -> list[Fact]:
    if path.suffix.lower() == ".json":
        data = json.loads(path.read_text(encoding="utf-8-sig"))
    else:
        data = yaml.safe_load(path.read_text(encoding="utf-8-sig"))
    facts = []
    for label, value, locator in _flatten(data):
        fact = _make_fact(path, locator, label, value)
        if fact:
            facts.append(fact)
    return facts


def _extract_csv(path: Path) -> list[Fact]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        return []
    facts: list[Fact] = []
    if all(len(row) == 2 for row in rows) and len(rows) > 1:
        for index, (label, value) in enumerate(rows, start=1):
            fact = _make_fact(path, f"row {index}", label, value)
            if fact:
                facts.append(fact)
        return facts
    headers = rows[0]
    for row_index, row in enumerate(rows[1:], start=2):
        for column_index, label in enumerate(headers):
            if column_index < len(row):
                fact = _make_fact(path, f"row {row_index}, column {label}", label, row[column_index])
                if fact:
                    facts.append(fact)
    return facts


def _facts_from_text(path: Path, text: str, locator_prefix: str = "line") -> list[Fact]:
    facts: list[Fact] = []
    seen: set[tuple[str, str]] = set()
    for number, line in enumerate(text.splitlines(), start=1):
        cleaned = re.sub(r"^[#>\s]+", "", line).strip()
        match = PAIR_PATTERN.match(cleaned)
        if match:
            label, value = match.groups()
            fact = _make_fact(path, f"{locator_prefix} {number}", label, value)
            if fact and (fact.key, fact.normalized_value) not in seen:
                facts.append(fact)
                seen.add((fact.key, fact.normalized_value))

        # High-precision facts can be recovered even from prose and resumes.
        for label, pattern in (("email", EMAIL_PATTERN), ("phone", PHONE_PATTERN)):
            for found in pattern.findall(cleaned):
                fact = _make_fact(path, f"{locator_prefix} {number}", label, found)
                if fact and (fact.key, fact.normalized_value) not in seen:
                    facts.append(fact)
                    seen.add((fact.key, fact.normalized_value))
        for found in URL_PATTERN.findall(cleaned):
            host_key = "linkedin" if "linkedin.com" in found else "github" if "github.com" in found else "website"
            fact = _make_fact(path, f"{locator_prefix} {number}", host_key, found.rstrip(".,"))
            if fact and (fact.key, fact.normalized_value) not in seen:
                facts.append(fact)
                seen.add((fact.key, fact.normalized_value))
    return facts


def _extract_docx(path: Path) -> list[Fact]:
    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                lines.append(f"{cells[0]}: {cells[1]}")
    return _facts_from_text(path, "\n".join(lines), "paragraph")


def _extract_pdf(path: Path) -> list[Fact]:
    reader = PdfReader(path)
    facts: list[Fact] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        facts.extend(_facts_from_text(path, text, f"page {page_number}, line"))
    return facts


def extract_file(path: str | Path) -> list[Fact]:
    source = Path(path).expanduser().resolve()
    if not source.exists():
        raise ExtractionError(f"Source does not exist: {source}")
    extension = source.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported source type '{extension}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    try:
        if extension in {".json", ".yaml", ".yml"}:
            return _extract_structured(source)
        if extension == ".csv":
            return _extract_csv(source)
        if extension in {".txt", ".md"}:
            return _facts_from_text(source, source.read_text(encoding="utf-8-sig"))
        if extension == ".docx":
            return _extract_docx(source)
        return _extract_pdf(source)
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Could not read {source.name}: {exc}") from exc


def extract_many(paths: list[str | Path]) -> list[Fact]:
    return [fact for path in paths for fact in extract_file(path)]
